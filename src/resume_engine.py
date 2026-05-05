"""
resume_engine.py — Murphy Resume Builder Engine (PATCH-193)

Capabilities:
  - Parse uploaded resume (PDF, DOCX, TXT, plain text)
  - Polish with LLM (rewrite bullets, fix grammar, strengthen impact language)
  - Tailor to a job description
  - Generate PDF output (reportlab — no system deps)
  - Forge-registered function: polish_resume(raw_text, job_description="") -> dict

DB: /var/lib/murphy-production/resume.db
Uploads: /var/lib/murphy-production/uploads/resumes/
"""

from __future__ import annotations
import os, re, json, uuid, sqlite3, logging, io, hashlib
from datetime import datetime, timezone
from typing import Optional, Dict, List
from pathlib import Path

logger = logging.getLogger(__name__)

RESUME_DB      = "/var/lib/murphy-production/resume.db"
RESUME_UPLOADS = "/var/lib/murphy-production/uploads/resumes"
RESUME_PDFS    = "/var/lib/murphy-production/uploads/resume_pdfs"

os.makedirs(RESUME_UPLOADS, exist_ok=True)
os.makedirs(RESUME_PDFS, exist_ok=True)


# ══════════════════════════════════════════════════════════════════════════════
# DATABASE
# ══════════════════════════════════════════════════════════════════════════════

def _db() -> sqlite3.Connection:
    conn = sqlite3.connect(RESUME_DB, timeout=8)
    conn.row_factory = sqlite3.Row
    return conn

def ensure_tables():
    with _db() as conn:
        conn.executescript("""
            CREATE TABLE IF NOT EXISTS resumes (
                id           TEXT PRIMARY KEY,
                user_id      TEXT DEFAULT 'default',
                filename     TEXT DEFAULT '',
                raw_text     TEXT DEFAULT '',
                parsed_json  TEXT DEFAULT '{}',
                polished_json TEXT DEFAULT '{}',
                job_desc     TEXT DEFAULT '',
                pdf_path     TEXT DEFAULT '',
                status       TEXT DEFAULT 'draft',
                created_at   TEXT NOT NULL,
                updated_at   TEXT NOT NULL
            );
            CREATE INDEX IF NOT EXISTS idx_resume_user ON resumes(user_id);
        """)
        conn.commit()


# ══════════════════════════════════════════════════════════════════════════════
# FILE PARSING
# ══════════════════════════════════════════════════════════════════════════════

def parse_uploaded_file(file_bytes: bytes, filename: str) -> str:
    """Extract plain text from PDF, DOCX, or TXT file bytes."""
    ext = Path(filename).suffix.lower()

    if ext == ".pdf":
        try:
            import pypdfium2 as pdfium
            doc = pdfium.PdfDocument(file_bytes)
            pages = []
            for i in range(len(doc)):
                page = doc[i]
                textpage = page.get_textpage()
                pages.append(textpage.get_text_range())
            return "\n".join(pages)
        except Exception as e:
            logger.warning("pypdfium2 failed: %s — trying pdfminer", e)
            try:
                from pdfminer.high_level import extract_text_to_fp
                from pdfminer.layout import LAParams
                buf = io.BytesIO(file_bytes)
                out = io.StringIO()
                extract_text_to_fp(buf, out, laparams=LAParams())
                return out.getvalue()
            except Exception as e2:
                logger.error("PDF parse failed: %s", e2)
                return ""

    elif ext in (".docx", ".doc"):
        try:
            import docx
            doc = docx.Document(io.BytesIO(file_bytes))
            return "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        except Exception as e:
            logger.error("DOCX parse failed: %s", e)
            return ""

    else:
        # TXT or plain paste
        try:
            return file_bytes.decode("utf-8", errors="replace")
        except Exception:
            return ""


# ══════════════════════════════════════════════════════════════════════════════
# RESUME SECTION PARSER
# ══════════════════════════════════════════════════════════════════════════════

SECTION_HEADERS = {
    "contact":     r"(?i)(contact|personal info|personal information)",
    "summary":     r"(?i)(summary|objective|profile|about me|professional summary)",
    "experience":  r"(?i)(experience|work history|employment|work experience|career)",
    "education":   r"(?i)(education|academic|qualifications|degrees?)",
    "skills":      r"(?i)(skills?|technical skills?|competencies|expertise|technologies)",
    "projects":    r"(?i)(projects?|portfolio|work samples?|accomplishments?)",
    "certifications": r"(?i)(certif|licenses?|credentials?|awards?)",
    "other":       r"(?i)(volunteer|interests?|hobbies|languages?|publications?)",
}

def parse_sections(text: str) -> Dict[str, str]:
    """Split resume text into named sections."""
    lines = text.split("\n")
    sections: Dict[str, List[str]] = {k: [] for k in SECTION_HEADERS}
    sections["header"] = []
    current = "header"

    for line in lines:
        stripped = line.strip()
        matched = False
        for sec, pattern in SECTION_HEADERS.items():
            if re.match(pattern, stripped) and len(stripped) < 60:
                current = sec
                matched = True
                break
        if not matched:
            sections[current].append(line)

    return {k: "\n".join(v).strip() for k, v in sections.items() if "\n".join(v).strip()}


def extract_contact(text: str) -> Dict[str, str]:
    """Pull name, email, phone, LinkedIn from raw text."""
    contact = {"name": "", "email": "", "phone": "", "linkedin": "", "location": ""}
    # Email
    m = re.search(r"[a-zA-Z0-9._%+\-]+@[a-zA-Z0-9.\-]+\.[a-zA-Z]{2,}", text)
    if m: contact["email"] = m.group(0)
    # Phone
    m = re.search(r"[\+]?[(]?[0-9]{3}[)]?[-\s\.]?[0-9]{3}[-\s\.]?[0-9]{4}", text)
    if m: contact["phone"] = m.group(0)
    # LinkedIn
    m = re.search(r"linkedin\.com/in/[\w\-]+", text, re.I)
    if m: contact["linkedin"] = "https://www." + m.group(0)
    # Name — first non-empty line that looks like a name
    for line in text.split("\n")[:8]:
        line = line.strip()
        if line and len(line) < 50 and re.match(r"^[A-Z][a-zA-Z\-\.\']+(\s+[A-Z][a-zA-Z\-\.\']+)+$", line):
            contact["name"] = line
            break
    return contact


# ══════════════════════════════════════════════════════════════════════════════
# LLM POLISH
# ══════════════════════════════════════════════════════════════════════════════

def _llm(prompt: str, max_tokens: int = 2000) -> str:
    try:
        from src.llm_provider import MurphyLLMProvider
        llm = MurphyLLMProvider()
        result = llm.complete(prompt, max_tokens=max_tokens, temperature=0.3)
        # Handle LLMCompletion object or plain string
        if hasattr(result, "text"):
            return str(result.text or "")
        if hasattr(result, "content"):
            return str(result.content or "")
        if hasattr(result, "__str__"):
            s = str(result)
            # Avoid returning repr of an object
            if not s.startswith("<"):
                return s
        return ""
    except Exception as e:
        logger.error("LLM failed: %s", e)
        return ""

POLISH_SYSTEM = """You are an expert resume writer and career coach.
Your job is to take raw resume content and transform it into polished, 
impactful, ATS-optimized resume content. Rules:
- Use strong action verbs (Led, Built, Drove, Increased, Reduced, etc.)
- Quantify achievements where possible (%, $, time saved, team size)
- Keep bullets concise — 1-2 lines max
- Remove filler words and passive voice
- Maintain truthfulness — never fabricate details
- Output clean plain text, section by section
- For each section, output: SECTION: <name> then the content"""

def polish_resume(raw_text: str, job_description: str = "") -> Dict:
    """
    FORGE FUNCTION: Polish a raw resume with LLM.
    This is the registered Forge function — callable via /api/forge/invoke.
    
    Args:
        raw_text: Raw resume text (pasted or extracted from file)
        job_description: Optional job description to tailor the resume to
    
    Returns:
        dict with keys: polished_text, sections, contact, score, suggestions
    """
    ensure_tables()

    sections = parse_sections(raw_text)
    contact  = extract_contact(raw_text)

    # Build tailoring context
    tailor_ctx = ""
    if job_description.strip():
        tailor_ctx = f"""

Additionally, tailor this resume to the following job description.
Highlight relevant skills and experience that match:

JOB DESCRIPTION:
{job_description[:1500]}
"""

    prompt = f"""{POLISH_SYSTEM}

Here is the raw resume to polish:{tailor_ctx}

RAW RESUME:
{raw_text[:4000]}

Polish each section and return clean, impactful content.
Start with CONTACT INFO then go through each section.
End with a SUGGESTIONS section listing 3-5 specific improvements the person could make."""

    polished_raw = _llm(prompt, max_tokens=2500)
    if not polished_raw:
        polished_raw = raw_text  # fallback — return original if LLM fails

    # Parse polished output into sections
    polished_sections: Dict[str, str] = {}
    current_sec = "summary"
    for line in polished_raw.split("\n"):
        m = re.match(r"^SECTION:\s*(.+)$", line.strip(), re.I)
        if m:
            current_sec = m.group(1).strip().lower()
            polished_sections[current_sec] = ""
        else:
            polished_sections.setdefault(current_sec, "")
            polished_sections[current_sec] += line + "\n"

    # Extract suggestions
    suggestions = []
    sugg_text = polished_sections.get("suggestions", "")
    for line in sugg_text.split("\n"):
        line = line.strip().lstrip("-•123456789. ")
        if line and len(line) > 10:
            suggestions.append(line)

    # Score the original resume
    score = _score_resume(raw_text, sections)

    return {
        "contact":          contact,
        "sections":         sections,
        "polished_text":    polished_raw,
        "polished_sections": polished_sections,
        "suggestions":      suggestions[:5],
        "score":            score,
        "tailored":         bool(job_description.strip()),
    }


def _score_resume(text: str, sections: Dict) -> Dict:
    """Score resume on key dimensions. Returns dict with scores 0-100."""
    score = {
        "overall": 0,
        "completeness": 0,
        "impact_language": 0,
        "length": 0,
        "ats_friendly": 0,
    }
    # Completeness — which key sections exist
    key_sections = ["experience", "education", "skills", "summary", "contact"]
    found = sum(1 for s in key_sections if sections.get(s, "").strip())
    score["completeness"] = int((found / len(key_sections)) * 100)

    # Impact language — action verbs
    action_verbs = ["led","built","drove","increased","reduced","managed","created",
                    "developed","launched","delivered","achieved","improved","designed",
                    "implemented","optimized","generated","saved","grew","scaled"]
    text_lower = text.lower()
    verb_count = sum(1 for v in action_verbs if v in text_lower)
    score["impact_language"] = min(100, verb_count * 12)

    # Length — 400-800 words is ideal
    word_count = len(text.split())
    if 400 <= word_count <= 800:
        score["length"] = 100
    elif word_count < 400:
        score["length"] = int((word_count / 400) * 100)
    else:
        score["length"] = max(50, 100 - int((word_count - 800) / 10))

    # ATS friendly — has email, consistent formatting
    has_email = bool(re.search(r"@", text))
    has_dates  = bool(re.search(r"\b(20\d\d|19\d\d)\b", text))
    score["ats_friendly"] = (50 if has_email else 0) + (50 if has_dates else 0)

    score["overall"] = int(sum([
        score["completeness"] * 0.3,
        score["impact_language"] * 0.3,
        score["length"] * 0.2,
        score["ats_friendly"] * 0.2,
    ]))
    return score


# ══════════════════════════════════════════════════════════════════════════════
# PDF GENERATOR
# ══════════════════════════════════════════════════════════════════════════════

def _resume_to_markdown(resume_data: Dict) -> str:
    """Convert polished resume data to clean Markdown for ExportPipeline."""
    contact  = resume_data.get("contact", {})
    polished = resume_data.get("polished_sections", {})
    raw_secs = resume_data.get("sections", {})

    def sec(name):
        return (polished.get(name) or raw_secs.get(name) or "").strip()

    lines = []

    # Header
    name = contact.get("name", "")
    if name:
        lines.append(f"# {name}")
    parts = []
    if contact.get("email"):    parts.append(contact["email"])
    if contact.get("phone"):    parts.append(contact["phone"])
    if contact.get("linkedin"): parts.append(contact["linkedin"])
    if contact.get("location"): parts.append(contact["location"])
    if parts:
        lines.append("  |  ".join(parts))
    lines.append("")
    lines.append("---")
    lines.append("")

    section_order = ["summary", "experience", "education", "skills",
                     "projects", "certifications", "other"]
    for sec_name in section_order:
        content = sec(sec_name)
        if not content:
            continue
        lines.append(f"## {sec_name.title()}")
        lines.append("")
        for line in content.split("\n"):
            stripped = line.strip()
            if not stripped:
                lines.append("")
            elif stripped.startswith(("-", "•", "*")):
                lines.append("- " + stripped.lstrip("-•* "))
            else:
                lines.append(stripped)
        lines.append("")

    return "\n".join(lines)


def generate_pdf(resume_data: Dict, output_path: str) -> str:
    """
    Generate a polished resume PDF using Murphy's existing ExportPipeline +
    RichPDFRenderer (WeasyPrint → reportlab fallback). No duplication.
    Returns the output path.
    """
    import base64, sys
    sys.path.insert(0, "/opt/Murphy-System")

    try:
        from src.document_export.export_pipeline import ExportPipeline
        from src.document_export.brand_registry import BrandProfile

        brand = BrandProfile(
            brand_id="resume_default",
            company_name="Murphy System",
            primary_color="#00d4aa",
            secondary_color="#1a1a2e",
            accent_color="#00d4aa",
            font_heading="Helvetica",
            font_body="Helvetica",
            legal_line="Generated by Murphy System — murphy.systems",
        )

        pipeline = ExportPipeline(brand_registry=None)
        markdown_text = _resume_to_markdown(resume_data)
        pdf_b64 = pipeline._markdown_to_pdf(markdown_text, brand)
        pdf_bytes = base64.b64decode(pdf_b64)

        with open(output_path, "wb") as f:
            f.write(pdf_bytes)
        logger.info("[Resume] PDF written via ExportPipeline: %s (%d bytes)", output_path, len(pdf_bytes))
        return output_path

    except Exception as e:
        logger.warning("[Resume] ExportPipeline PDF failed (%s), using reportlab direct", e)
        # Reportlab direct fallback
        return _generate_pdf_reportlab(resume_data, output_path)


def _generate_pdf_reportlab(resume_data: Dict, output_path: str) -> str:
    """Direct reportlab fallback — minimal but functional."""
    from reportlab.lib.pagesizes import LETTER
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.lib import colors
    from reportlab.platypus import (SimpleDocTemplate, Paragraph, Spacer,
                                     HRFlowable, KeepTogether)
    from reportlab.lib.enums import TA_CENTER, TA_LEFT

    doc = SimpleDocTemplate(
        output_path,
        pagesize=LETTER,
        leftMargin=0.75*inch, rightMargin=0.75*inch,
        topMargin=0.6*inch, bottomMargin=0.6*inch,
    )

    # ── Styles ──────────────────────────────────────────────────────────────
    TEAL   = colors.HexColor("#00d4aa")
    DARK   = colors.HexColor("#1a1a2e")
    GRAY   = colors.HexColor("#555555")
    LGRAY  = colors.HexColor("#888888")

    styles = getSampleStyleSheet()

    name_style = ParagraphStyle("Name", fontSize=22, fontName="Helvetica-Bold",
                                 textColor=DARK, spaceAfter=2, alignment=TA_CENTER)
    contact_style = ParagraphStyle("Contact", fontSize=9, fontName="Helvetica",
                                    textColor=GRAY, spaceAfter=6, alignment=TA_CENTER)
    section_style = ParagraphStyle("Section", fontSize=11, fontName="Helvetica-Bold",
                                    textColor=TEAL, spaceBefore=10, spaceAfter=3)
    body_style = ParagraphStyle("Body", fontSize=9.5, fontName="Helvetica",
                                 textColor=DARK, spaceAfter=4, leading=13)
    bullet_style = ParagraphStyle("Bullet", fontSize=9.5, fontName="Helvetica",
                                   textColor=DARK, spaceAfter=3, leading=13,
                                   leftIndent=14, bulletIndent=4)
    job_title_style = ParagraphStyle("JobTitle", fontSize=10.5, fontName="Helvetica-Bold",
                                      textColor=DARK, spaceAfter=1, spaceBefore=6)
    date_style = ParagraphStyle("Date", fontSize=9, fontName="Helvetica-Oblique",
                                 textColor=LGRAY, spaceAfter=3)

    story = []
    contact  = resume_data.get("contact", {})
    polished = resume_data.get("polished_sections", {})
    raw_secs = resume_data.get("sections", {})

    # Use polished if available, fallback to raw
    def sec(name):
        return polished.get(name, raw_secs.get(name, "")).strip()

    # ── Header ───────────────────────────────────────────────────────────────
    name = contact.get("name", "") or resume_data.get("name", "Your Name")
    story.append(Paragraph(name, name_style))

    contact_parts = []
    if contact.get("email"):    contact_parts.append(contact["email"])
    if contact.get("phone"):    contact_parts.append(contact["phone"])
    if contact.get("linkedin"): contact_parts.append(contact["linkedin"])
    if contact.get("location"): contact_parts.append(contact["location"])
    if contact_parts:
        story.append(Paragraph("  •  ".join(contact_parts), contact_style))

    story.append(HRFlowable(width="100%", thickness=2, color=TEAL, spaceAfter=6))

    # ── Sections ─────────────────────────────────────────────────────────────
    section_order = ["summary", "experience", "education", "skills",
                     "projects", "certifications", "other"]

    for sec_name in section_order:
        content = sec(sec_name)
        if not content:
            continue

        story.append(Paragraph(sec_name.upper(), section_style))
        story.append(HRFlowable(width="100%", thickness=0.5, color=TEAL, spaceAfter=4))

        for line in content.split("\n"):
            line = line.strip()
            if not line:
                story.append(Spacer(1, 3))
                continue
            # Bullet points
            if line.startswith(("-", "•", "*", "·")):
                clean = line.lstrip("-•*· ").strip()
                story.append(Paragraph("• " + clean, bullet_style))
            # Date ranges
            elif re.search(r"\b(20\d\d|19\d\d)\b", line) and len(line) < 60:
                story.append(Paragraph(line, date_style))
            # Short bold-worthy lines (job titles / company names)
            elif len(line) < 70 and line[0].isupper() and not line.endswith("."):
                story.append(Paragraph("<b>" + line + "</b>", job_title_style))
            else:
                story.append(Paragraph(line, body_style))

    # ── Footer ───────────────────────────────────────────────────────────────
    story.append(Spacer(1, 12))
    story.append(HRFlowable(width="100%", thickness=0.5, color=TEAL))
    story.append(Paragraph(
        f"<font color='#888888' size='7'>Generated by Murphy System — murphy.systems — {datetime.now().strftime('%B %Y')}</font>",
        ParagraphStyle("Footer", alignment=TA_CENTER, spaceAfter=0)
    ))

    doc.build(story)
    return output_path


# ══════════════════════════════════════════════════════════════════════════════
# FULL PIPELINE
# ══════════════════════════════════════════════════════════════════════════════

def build_resume(
    raw_text: str = "",
    file_bytes: bytes = b"",
    filename: str = "",
    job_description: str = "",
    user_id: str = "default",
    name_override: str = "",
) -> Dict:
    """
    Full pipeline: parse → polish → generate PDF.
    Returns: {id, pdf_url, pdf_path, polished, score, suggestions, contact}
    """
    ensure_tables()

    # 1. Extract text
    if file_bytes and filename:
        text = parse_uploaded_file(file_bytes, filename)
    else:
        text = raw_text

    if not text.strip():
        return {"error": "No resume content provided", "success": False}

    # 2. Polish with LLM
    polished = polish_resume(text, job_description)
    if name_override:
        polished["contact"]["name"] = name_override

    # 3. Generate PDF
    rid      = str(uuid.uuid4())[:8]
    pdf_path = os.path.join(RESUME_PDFS, f"resume_{rid}.pdf")
    try:
        generate_pdf(polished, pdf_path)
        pdf_ok = True
    except Exception as e:
        logger.error("PDF generation failed: %s", e)
        pdf_ok  = False
        pdf_path = ""

    # 4. Store in DB
    now = datetime.now(timezone.utc).isoformat()
    with _db() as conn:
        conn.execute("""
            INSERT INTO resumes
            (id, user_id, filename, raw_text, parsed_json, polished_json,
             job_desc, pdf_path, status, created_at, updated_at)
            VALUES (?,?,?,?,?,?,?,?,?,?,?)
        """, (
            rid, user_id, filename or "pasted",
            text[:5000], json.dumps(polished.get("sections", {})),
            json.dumps(polished), job_description[:2000],
            pdf_path, "complete" if pdf_ok else "polished",
            now, now
        ))
        conn.commit()

    pdf_url = f"/api/resume/download/{rid}" if pdf_ok else None

    return {
        "success":     True,
        "id":          rid,
        "pdf_url":     pdf_url,
        "pdf_path":    pdf_path,
        "contact":     polished.get("contact", {}),
        "score":       polished.get("score", {}),
        "suggestions": polished.get("suggestions", []),
        "polished_text": polished.get("polished_text", ""),
        "polished_sections": polished.get("polished_sections", {}),
        "tailored":    bool(job_description.strip()),
    }
