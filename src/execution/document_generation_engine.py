"""
Document Generation Engine - Generate working documents (PDF, Word, HTML)
"""

import json
import logging
import uuid
from datetime import datetime, timezone
from enum import Enum
from html import escape as html_escape
from typing import Any, Dict, List, Optional

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


class DocumentType(Enum):
    """Types of documents"""
    PDF = "pdf"
    WORD = "word"
    HTML = "html"
    MARKDOWN = "markdown"
    PLAIN_TEXT = "plain_text"


class DocumentTemplate:
    """Document template"""

    def __init__(
        self,
        template_id: str,
        template_type: DocumentType,
        content: str,
        placeholders: Optional[List[str]] = None,
        styling: Optional[Dict] = None
    ):
        self.template_id = template_id
        self.template_type = template_type
        self.content = content
        self.placeholders = placeholders or []
        self.styling = styling or {}

    def render(self, context: Dict) -> str:
        """Render template with context"""
        rendered_content = self.content

        for placeholder in self.placeholders:
            placeholder_key = f"{{{placeholder}}}"
            value = str(context.get(placeholder, ""))
            rendered_content = rendered_content.replace(placeholder_key, value)

        return rendered_content


class Document:
    """Generated document"""

    def __init__(
        self,
        document_id: Optional[str] = None,
        document_type: DocumentType = DocumentType.PDF,
        content: str = "",
        metadata: Optional[Dict] = None
    ):
        self.document_id = document_id or str(uuid.uuid4())
        self.document_type = document_type
        self.content = content
        self.metadata = metadata or {}
        self.created_at = datetime.now(timezone.utc)

    def to_dict(self) -> Dict:
        """Convert document to dictionary"""
        return {
            'document_id': self.document_id,
            'document_type': self.document_type.value,
            'content': self.content,
            'metadata': self.metadata,
            'created_at': self.created_at.isoformat()
        }


class DocumentGenerationEngine:
    """Generate documents from templates"""

    def __init__(self):
        self.templates: Dict[str, DocumentTemplate] = {}
        self.documents: Dict[str, Document] = {}

    def register_template(self, template: DocumentTemplate) -> None:
        """Register a document template"""
        self.templates[template.template_id] = template
        logger.info(f"Template registered: {template.template_id}")

    def generate_pdf(
        self,
        template_id: str,
        data: Dict,
        metadata: Optional[Dict] = None
    ) -> Document:
        """Generate PDF document"""
        template = self.templates.get(template_id)
        if not template:
            raise ValueError(f"Template not found: {template_id}")

        # Render template
        content = template.render(data)

        # Convert to PDF format
        pdf_content = self._convert_to_pdf(content, template.styling)

        document = Document(
            document_type=DocumentType.PDF,
            content=pdf_content,
            metadata=metadata
        )

        self.documents[document.document_id] = document
        logger.info(f"PDF document generated: {document.document_id}")

        return document

    def generate_word(
        self,
        template_id: str,
        data: Dict,
        metadata: Optional[Dict] = None
    ) -> Document:
        """Generate Word document"""
        template = self.templates.get(template_id)
        if not template:
            raise ValueError(f"Template not found: {template_id}")

        # Render template
        content = template.render(data)

        # Convert to Word format
        word_content = self._convert_to_word(content, template.styling)

        document = Document(
            document_type=DocumentType.WORD,
            content=word_content,
            metadata=metadata
        )

        self.documents[document.document_id] = document
        logger.info(f"Word document generated: {document.document_id}")

        return document

    def generate_html(
        self,
        template_id: str,
        data: Dict,
        metadata: Optional[Dict] = None
    ) -> Document:
        """Generate HTML document"""
        template = self.templates.get(template_id)
        if not template:
            raise ValueError(f"Template not found: {template_id}")

        # Render template
        content = template.render(data)

        # Wrap in HTML structure
        html_content = self._wrap_in_html(content, template.styling)

        document = Document(
            document_type=DocumentType.HTML,
            content=html_content,
            metadata=metadata
        )

        self.documents[document.document_id] = document
        logger.info(f"HTML document generated: {document.document_id}")

        return document

    def generate_from_template(
        self,
        template_id: str,
        context: Dict,
        metadata: Optional[Dict] = None
    ) -> Document:
        """Generate document from template"""
        template = self.templates.get(template_id)
        if not template:
            raise ValueError(f"Template not found: {template_id}")

        # Render template
        content = template.render(context)

        document = Document(
            document_type=template.template_type,
            content=content,
            metadata=metadata
        )

        self.documents[document.document_id] = document
        logger.info(f"Document generated: {document.document_id}")

        return document

    def preview_document(
        self,
        template_id: str,
        data: Dict
    ) -> str:
        """Preview document content"""
        template = self.templates.get(template_id)
        if not template:
            raise ValueError(f"Template not found: {template_id}")

        return template.render(data)

    def _convert_to_pdf(self, content: str, styling: Dict) -> str:
        """Convert content to PDF format.

        Tries WeasyPrint → reportlab → text fallback.
        Returns base64-encoded PDF data when a PDF library is available.
        """
        # 1. Try WeasyPrint for rich rendering
        try:
            from document_export.brand_registry import BrandProfile
            from document_export.pdf_renderer import RichPDFRenderer

            renderer = RichPDFRenderer()
            if renderer.is_available():
                brand = BrandProfile()
                wrapped_html = (
                    "<html><body>"
                    f"<pre style='font-family: monospace; white-space: pre-wrap;'>"
                    f"{html_escape(content)}"
                    "</pre></body></html>"
                )
                return renderer.render_to_base64(wrapped_html, brand, {})
        except Exception as exc:  # noqa: BLE001
            logger.debug("Rich PDF renderer failed, trying fallback: %s", exc)

        # 2. reportlab fallback
        try:
            import base64
            from io import BytesIO

            from reportlab.lib.pagesizes import letter  # type: ignore[import-untyped]
            from reportlab.pdfgen import canvas as rl_canvas  # type: ignore[import-untyped]

            buf = BytesIO()
            c = rl_canvas.Canvas(buf, pagesize=letter)
            font_name = styling.get("font", "Helvetica")
            font_size = int(styling.get("size", 12))
            c.setFont(font_name, font_size)

            # Simple text flow
            y = 750
            for line in content.split("\n"):
                if y < 50:
                    c.showPage()
                    c.setFont(font_name, font_size)
                    y = 750
                c.drawString(72, y, line)
                y -= font_size + 2
            c.save()
            return base64.b64encode(buf.getvalue()).decode("ascii")
        except ImportError:
            # Graceful fallback — return structured text
            return f"%PDF-1.4-TEXT-FALLBACK\n{content}\n%%EOF"

    def _convert_to_word(self, content: str, styling: Dict) -> str:
        """Convert content to Word (OOXML) format.

        Uses python-docx when available; otherwise returns a minimal
        Office Open XML document string.
        Returns base64-encoded DOCX data when python-docx is available.
        """
        try:
            import base64
            from io import BytesIO

            from docx import Document as DocxDocument  # type: ignore[import-untyped]
            from docx.shared import Pt  # type: ignore[import-untyped]

            doc = DocxDocument()
            font_name = styling.get("font", "Calibri")
            font_size = int(styling.get("size", 12))
            for para_text in content.split("\n"):
                p = doc.add_paragraph(para_text)
                for run in p.runs:
                    run.font.name = font_name
                    run.font.size = Pt(font_size)
            buf = BytesIO()
            doc.save(buf)
            return base64.b64encode(buf.getvalue()).decode("ascii")
        except ImportError:
            # Graceful fallback — minimal XML payload
            escaped = content.replace("&", "&amp;").replace("<", "&lt;")
            return (
                '<?xml version="1.0" encoding="UTF-8"?>\n'
                "<document>\n"
                f"  <body>{escaped}</body>\n"
                "</document>"
            )

    def _wrap_in_html(self, content: str, styling: Dict) -> str:
        """Wrap content in HTML structure"""
        styling_str = ""
        if styling:
            for key, value in styling.items():
                styling_str += f"{key}: {value}; "

        html_content = f"""<!DOCTYPE html>
<html>
<head>
    <style>
        body {{ {styling_str} }}
    </style>
</head>
<body>
    {content}
</body>
</html>"""

        return html_content

    def render_pdf(self, content: str, styling: Optional[Dict] = None) -> str:
        """Render plain-text *content* as a PDF document.

        This is the public API for callers that have pre-rendered content and
        just need the format conversion.  Delegates to :meth:`_convert_to_pdf`.

        Returns base64-encoded PDF bytes (or a text-fallback sentinel when
        reportlab is not installed).
        """
        return self._convert_to_pdf(content, styling or {})

    def render_word(self, content: str, styling: Optional[Dict] = None) -> str:
        """Render plain-text *content* as a Word (.docx) document.

        This is the public API for callers that have pre-rendered content and
        just need the format conversion.  Delegates to :meth:`_convert_to_word`.

        Returns base64-encoded DOCX bytes (or a minimal XML fallback when
        python-docx is not installed).
        """
        return self._convert_to_word(content, styling or {})

    def get_document(self, document_id: str) -> Optional[Document]:
        """Get document by ID"""
        return self.documents.get(document_id)

    def get_all_documents(self) -> List[Dict]:
        """Get all documents"""
        return [doc.to_dict() for doc in self.documents.values()]


# Convenience functions

def create_template(
    template_id: str,
    template_type: DocumentType,
    content: str,
    placeholders: Optional[List[str]] = None
) -> DocumentTemplate:
    """Create a document template"""
    return DocumentTemplate(
        template_id=template_id,
        template_type=template_type,
        content=content,
        placeholders=placeholders
    )


def generate_document(
    engine: DocumentGenerationEngine,
    template_id: str,
    data: Dict,
    metadata: Optional[Dict] = None
) -> Dict:
    """Generate a document"""
    document = engine.generate_from_template(template_id, data, metadata)
    return document.to_dict()
