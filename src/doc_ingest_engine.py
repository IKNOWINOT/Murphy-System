"""
Document Ingestion Engine — PATCH-080a
Ingests engineering documents into Murphy's knowledge base.

Supported formats:
  PDF   — technical specs, drawings, reports, RFIs
  DOCX  — specifications, scopes of work, submittals
  XLSX  — estimating spreadsheets, takeoff sheets, bid tabs
  DXF   — CAD drawings (mechanical, electrical, plumbing)
  TXT   — raw text specs

All ingested documents are:
  1. Parsed and chunked
  2. Metadata-tagged (discipline, drawing_number, revision, etc.)
  3. Stored in Murphy's Knowledge Graph
  4. Indexed for RAG (retrieval-augmented generation)

PATCH-080a | Label: DOC-INGEST-001
Copyright © 2020-2026 Inoni LLC
"""
from __future__ import annotations

import hashlib
import json
import logging
import os
import re
import time
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any, Dict, List, Optional

logger = logging.getLogger(__name__)

UPLOAD_DIR = Path(os.environ.get("MURPHY_UPLOAD_DIR", "/var/lib/murphy-production/uploads"))
UPLOAD_DIR.mkdir(parents=True, exist_ok=True)


@dataclass
class DocumentChunk:
    """A chunk of extracted document content."""
    doc_id: str
    chunk_index: int
    text: str
    metadata: Dict[str, Any]
    source_file: str
    page: Optional[int] = None


@dataclass
class IngestedDocument:
    """Result of ingesting a document."""
    doc_id: str
    filename: str
    file_type: str
    discipline: str          # mechanical | electrical | plumbing | civil | structural | general
    drawing_number: str
    revision: str
    title: str
    chunks: List[DocumentChunk]
    tables: List[Dict]       # structured tables (from XLSX/PDF)
    metadata: Dict[str, Any]
    ok: bool
    error: str = ""


def _detect_discipline(text: str, filename: str) -> str:
    """Detect MEP discipline from content and filename."""
    combined = (text[:2000] + filename).lower()
    if any(k in combined for k in ["mechanical", "hvac", "ductwork", "mech", "-m-", "_m_", "M-"]):
        return "mechanical"
    if any(k in combined for k in ["electrical", "panel", "conduit", "elec", "-e-", "_e_", "E-"]):
        return "electrical"
    if any(k in combined for k in ["plumbing", "drain", "sanitary", "plumb", "-p-", "_p_", "P-"]):
        return "plumbing"
    if any(k in combined for k in ["structural", "steel", "concrete", "beam", "column", "-s-", "S-"]):
        return "structural"
    if any(k in combined for k in ["civil", "grading", "sitework", "utilities", "-c-", "C-"]):
        return "civil"
    if any(k in combined for k in ["architecture", "arch", "floor plan", "-a-", "A-"]):
        return "architectural"
    return "general"


def _extract_drawing_meta(text: str, filename: str) -> Dict[str, str]:
    """Extract drawing number, revision, title from text."""
    meta = {"drawing_number": "", "revision": "", "title": ""}
    # Drawing number patterns: A-101, M-201, E-301, DWG-001
    m = re.search(r"([A-Z]{1,2}[-_]\d{3,4}[A-Za-z]?)", text[:500] + filename)
    if m:
        meta["drawing_number"] = m.group(1)
    # Revision
    m2 = re.search(r"\bRev(?:ision)?[.:_\s]*([A-Z0-9]+)", text[:500], re.I)
    if m2:
        meta["revision"] = m2.group(1)
    # Title from first meaningful line
    for line in text.splitlines()[:20]:
        line = line.strip()
        if len(line) > 10 and not re.match(r"^[0-9/\-]+$", line):
            meta["title"] = line[:120]
            break
    return meta


def _chunk_text(text: str, doc_id: str, source_file: str, 
                metadata: Dict, chunk_size: int = 800) -> List[DocumentChunk]:
    """Split text into overlapping chunks for RAG."""
    words = text.split()
    chunks = []
    step = chunk_size - 100  # 100-word overlap
    for i, start in enumerate(range(0, len(words), step)):
        chunk_text = " ".join(words[start:start + chunk_size])
        if len(chunk_text.strip()) < 30:
            continue
        chunks.append(DocumentChunk(
            doc_id=doc_id,
            chunk_index=i,
            text=chunk_text,
            metadata=metadata,
            source_file=source_file,
        ))
    return chunks


def ingest_pdf(file_path: str) -> IngestedDocument:
    """Extract text + tables from PDF (technical specs, drawings, reports)."""
    doc_id = hashlib.sha256(Path(file_path).read_bytes()).hexdigest()[:16]
    fname = Path(file_path).name
    try:
        import pdfplumber
        text_parts = []
        tables = []
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                pt = page.extract_text() or ""
                text_parts.append(pt)
                for tbl in page.extract_tables():
                    if tbl:
                        tables.append({"page": page.page_number, "data": tbl})
        full_text = "\n".join(text_parts)
        discipline = _detect_discipline(full_text, fname)
        dmeta = _extract_drawing_meta(full_text, fname)
        chunks = _chunk_text(full_text, doc_id, fname,
                             {"discipline": discipline, **dmeta, "source": fname})
        logger.info("DOC-INGEST: PDF %s → %d chunks, %d tables, discipline=%s",
                    fname, len(chunks), len(tables), discipline)
        return IngestedDocument(
            doc_id=doc_id, filename=fname, file_type="pdf",
            discipline=discipline, chunks=chunks, tables=tables,
            metadata={"pages": len(text_parts), **dmeta}, ok=True,
            drawing_number=dmeta["drawing_number"],
            revision=dmeta["revision"], title=dmeta["title"],
        )
    except Exception as exc:
        logger.error("DOC-INGEST: PDF %s failed: %s", fname, exc)
        return IngestedDocument(doc_id=doc_id, filename=fname, file_type="pdf",
                                discipline="general", chunks=[], tables=[],
                                metadata={}, ok=False, error=str(exc),
                                drawing_number="", revision="", title="")


def ingest_docx(file_path: str) -> IngestedDocument:
    """Extract text from Word documents (specs, scopes of work)."""
    doc_id = hashlib.sha256(Path(file_path).read_bytes()).hexdigest()[:16]
    fname = Path(file_path).name
    try:
        import docx
        doc = docx.Document(file_path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        full_text = "\n".join(paragraphs)
        # Also extract table text
        tables = []
        for tbl in doc.tables:
            rows = [[cell.text.strip() for cell in row.cells] for row in tbl.rows]
            tables.append({"data": rows})
        discipline = _detect_discipline(full_text, fname)
        dmeta = _extract_drawing_meta(full_text, fname)
        chunks = _chunk_text(full_text, doc_id, fname,
                             {"discipline": discipline, **dmeta, "source": fname})
        logger.info("DOC-INGEST: DOCX %s → %d chunks, discipline=%s", fname, len(chunks), discipline)
        return IngestedDocument(
            doc_id=doc_id, filename=fname, file_type="docx",
            discipline=discipline, chunks=chunks, tables=tables,
            metadata={"paragraphs": len(paragraphs), **dmeta}, ok=True,
            drawing_number=dmeta["drawing_number"],
            revision=dmeta["revision"], title=dmeta["title"],
        )
    except Exception as exc:
        logger.error("DOC-INGEST: DOCX %s failed: %s", fname, exc)
        return IngestedDocument(doc_id=doc_id, filename=fname, file_type="docx",
                                discipline="general", chunks=[], tables=[],
                                metadata={}, ok=False, error=str(exc),
                                drawing_number="", revision="", title="")


def ingest_xlsx(file_path: str) -> IngestedDocument:
    """Extract estimating spreadsheets, bid tabs, takeoff sheets."""
    doc_id = hashlib.sha256(Path(file_path).read_bytes()).hexdigest()[:16]
    fname = Path(file_path).name
    try:
        import openpyxl
        wb = openpyxl.load_workbook(file_path, data_only=True)
        tables = []
        all_text = []
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            rows = []
            for row in ws.iter_rows(values_only=True):
                row_data = [str(c) if c is not None else "" for c in row]
                if any(r.strip() for r in row_data):
                    rows.append(row_data)
                    all_text.append(" ".join(r for r in row_data if r.strip()))
            if rows:
                tables.append({"sheet": sheet_name, "data": rows})
        full_text = "\n".join(all_text)
        discipline = _detect_discipline(full_text, fname)
        chunks = _chunk_text(full_text, doc_id, fname,
                             {"discipline": discipline, "source": fname, "type": "spreadsheet"})
        logger.info("DOC-INGEST: XLSX %s → %d sheets, %d chunks", fname, len(wb.sheetnames), len(chunks))
        return IngestedDocument(
            doc_id=doc_id, filename=fname, file_type="xlsx",
            discipline=discipline, chunks=chunks, tables=tables,
            metadata={"sheets": wb.sheetnames}, ok=True,
            drawing_number="", revision="", title=fname,
        )
    except Exception as exc:
        logger.error("DOC-INGEST: XLSX %s failed: %s", fname, exc)
        return IngestedDocument(doc_id=doc_id, filename=fname, file_type="xlsx",
                                discipline="general", chunks=[], tables=[],
                                metadata={}, ok=False, error=str(exc),
                                drawing_number="", revision="", title="")


def ingest_dxf(file_path: str) -> IngestedDocument:
    """Parse DXF CAD drawings — extract entities, text, layers, blocks."""
    doc_id = hashlib.sha256(Path(file_path).read_bytes()).hexdigest()[:16]
    fname = Path(file_path).name
    try:
        import ezdxf
        dxf_doc = ezdxf.readfile(file_path)
        msp = dxf_doc.modelspace()
        
        text_entities = []
        layers = set()
        blocks = []
        entity_counts: Dict[str, int] = {}
        
        for entity in msp:
            etype = entity.dxftype()
            entity_counts[etype] = entity_counts.get(etype, 0) + 1
            if hasattr(entity.dxf, "layer"):
                layers.add(entity.dxf.layer)
            if etype in ("TEXT", "MTEXT"):
                try:
                    txt = entity.dxf.text if etype == "TEXT" else entity.text
                    if txt and txt.strip():
                        text_entities.append(txt.strip())
                except Exception:
                    pass

        for block in dxf_doc.blocks:
            if not block.name.startswith("*"):
                blocks.append(block.name)

        full_text = " ".join(text_entities)
        layers_list = sorted(layers)
        discipline = _detect_discipline(full_text + " " + " ".join(layers_list), fname)
        dmeta = _extract_drawing_meta(full_text, fname)
        chunks = _chunk_text(full_text, doc_id, fname,
                             {"discipline": discipline, **dmeta,
                              "layers": layers_list[:20], "source": fname})
        
        logger.info("DOC-INGEST: DXF %s → %d layers, %d text entities, discipline=%s",
                    fname, len(layers), len(text_entities), discipline)
        return IngestedDocument(
            doc_id=doc_id, filename=fname, file_type="dxf",
            discipline=discipline, chunks=chunks, tables=[],
            metadata={
                "layers": layers_list, "blocks": blocks[:50],
                "entity_counts": entity_counts, **dmeta,
            },
            ok=True,
            drawing_number=dmeta["drawing_number"],
            revision=dmeta["revision"], title=dmeta["title"] or fname,
        )
    except Exception as exc:
        logger.error("DOC-INGEST: DXF %s failed: %s", fname, exc)
        return IngestedDocument(doc_id=doc_id, filename=fname, file_type="dxf",
                                discipline="general", chunks=[], tables=[],
                                metadata={}, ok=False, error=str(exc),
                                drawing_number="", revision="", title="")


def ingest_file(file_path: str) -> IngestedDocument:
    """Auto-detect format and ingest any supported engineering document."""
    ext = Path(file_path).suffix.lower()
    dispatch = {
        ".pdf": ingest_pdf,
        ".docx": ingest_docx,
        ".doc": ingest_docx,
        ".xlsx": ingest_xlsx,
        ".xls": ingest_xlsx,
        ".dxf": ingest_dxf,
        ".dwg": ingest_dxf,  # ezdxf can handle some DWG
    }
    handler = dispatch.get(ext)
    if handler is None:
        return IngestedDocument(
            doc_id="", filename=Path(file_path).name, file_type=ext,
            discipline="general", chunks=[], tables=[], metadata={},
            ok=False, error=f"Unsupported format: {ext}",
            drawing_number="", revision="", title="",
        )
    return handler(file_path)


def store_in_knowledge_graph(doc: IngestedDocument) -> bool:
    """Push ingested document chunks into Murphy's Knowledge Graph."""
    try:
        from src.murphy_memory_palace import MemoryPalace
        palace = MemoryPalace()
        stored = 0
        for chunk in doc.chunks:
            result = palace.index_conversation(
                text=chunk.text,
                source=doc.filename,
                metadata={
                    "doc_id": doc.doc_id,
                    "discipline": doc.discipline,
                    "drawing_number": doc.drawing_number,
                    "revision": doc.revision,
                    "chunk_index": chunk.chunk_index,
                    "file_type": doc.file_type,
                    **{k: v for k, v in doc.metadata.items() if isinstance(v, (str, int, float, bool))},
                },
            )
            if result.get("status") == "indexed":
                stored += 1
        logger.info("DOC-INGEST: %d/%d chunks stored in KG for %s",
                    stored, len(doc.chunks), doc.filename)
        return stored > 0
    except Exception as exc:
        logger.error("DOC-INGEST: KG store failed for %s: %s", doc.filename, exc)
        return False
